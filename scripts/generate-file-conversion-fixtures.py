#!/usr/bin/env python3
"""Generate reproducible, non-sensitive Zero File conversion fixtures."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import subprocess
import uuid
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as ReportLabImage,
    PageBreak,
    PageTemplate,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "tests" / "fixtures" / "fileConversion"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
INK = RGBColor(0x18, 0x1A, 0x1F)
FIXTURE_NAMES = (
    "rich-layout.docx",
    "large-structured.docx",
    "malformed.docx",
    "~$office-lock.docx",
    "rich-layout.pdf",
    "image-only-scan.pdf",
    "encrypted.pdf",
    "large-structured.pdf",
    "malformed.pdf",
    "unsupported.txt",
    "expected.json",
)
PDF_FONT_NAME = "ZeroFixtureCJK"
PDF_FONT_PATH: Path | None = None
DOCX_CJK_FONT_NAME = "Noto Sans SC"
RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
FONT_RELATIONSHIP_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"
)


def set_run_font(run, name: str = "Calibri", size: float = 11, *, bold: bool = False) -> None:
    effective_name = (
        DOCX_CJK_FONT_NAME
        if any("\u3400" <= character <= "\u9fff" for character in run.text)
        else name
    )
    run.font.name = effective_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), effective_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), effective_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), DOCX_CJK_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), effective_name)
    language = run._element.get_or_add_rPr().find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        run._element.get_or_add_rPr().append(language)
    language.set(qn("w:eastAsia"), "zh-CN")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = INK


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_CJK_FONT_NAME)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 12, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_CJK_FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("Zero File conversion fixture"), size=9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("Generated test content - no user data"), size=8)


def set_table_geometry(table, widths_inches: list[float]) -> None:
    table.autofit = False
    total_dxa = round(sum(widths_inches) * 1440)
    table_props = table._tbl.tblPr
    table_width = table_props.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_props.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(total_dxa))
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    table_props.append(indent)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(round(width * 1440)))


def subset_cjk_font(source: Path, target: Path) -> None:
    subsetter = shutil.which("hb-subset")
    if subsetter is None:
        raise RuntimeError("hb-subset is required to create the embedded DOCX fixture font")
    characters = OUTPUT / "assets" / "font-characters.txt"
    subprocess.run(
        [
            subsetter,
            str(source),
            f"--text-file={characters}",
            f"--output-file={target}",
            "--name-IDs=*",
            "--name-languages=*",
            "--layout-features=*",
        ],
        check=True,
        capture_output=True,
        text=True,
    )


def obfuscate_openxml_font(font_data: bytes, font_key: uuid.UUID) -> bytes:
    obfuscated = bytearray(font_data)
    key = font_key.bytes
    for index in range(min(32, len(obfuscated))):
        obfuscated[index] ^= key[15 - (index % 16)]
    return bytes(obfuscated)


def embed_docx_font(path: Path, font_path: Path) -> None:
    with zipfile.ZipFile(path) as package:
        entries = {name: package.read(name) for name in package.namelist()}

    font_data = font_path.read_bytes()
    font_key = uuid.UUID(bytes=hashlib.sha256(font_data).digest()[:16])
    font_key_value = "{" + str(font_key).upper() + "}"
    font_part = f"word/fonts/{DOCX_CJK_FONT_NAME}.odttf"
    relationship_id = "rId1"

    font_table = etree.fromstring(entries["word/fontTable.xml"])
    matching_fonts = font_table.xpath(
        "./w:font[@w:name=$name]",
        namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        name=DOCX_CJK_FONT_NAME,
    )
    font = matching_fonts[0] if matching_fonts else etree.SubElement(font_table, qn("w:font"))
    font.set(qn("w:name"), DOCX_CJK_FONT_NAME)
    for child_name, value in (
        ("w:charset", "86"),
        ("w:family", "auto"),
        ("w:pitch", "variable"),
    ):
        child = font.find(qn(child_name))
        if child is None:
            child = etree.SubElement(font, qn(child_name))
        child.set(qn("w:val"), value)
    signature = font.find(qn("w:sig"))
    if signature is None:
        signature = etree.SubElement(font, qn("w:sig"))
    for attribute, value in (
        ("w:usb0", "E0002AFF"),
        ("w:usb1", "C000247B"),
        ("w:usb2", "00000009"),
        ("w:usb3", "00000000"),
        ("w:csb0", "000001FF"),
        ("w:csb1", "00000000"),
    ):
        signature.set(qn(attribute), value)
    embed = font.find(qn("w:embedRegular"))
    if embed is None:
        embed = etree.SubElement(font, qn("w:embedRegular"))
    embed.set(qn("r:id"), relationship_id)
    embed.set(qn("w:fontKey"), font_key_value)
    entries["word/fontTable.xml"] = etree.tostring(
        font_table, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    relationships_path = "word/_rels/fontTable.xml.rels"
    if relationships_path in entries:
        relationships = etree.fromstring(entries[relationships_path])
    else:
        relationships = etree.Element(
            f"{{{RELATIONSHIPS_NS}}}Relationships", nsmap={None: RELATIONSHIPS_NS}
        )
    relationship = etree.SubElement(relationships, f"{{{RELATIONSHIPS_NS}}}Relationship")
    relationship.set("Id", relationship_id)
    relationship.set("Type", FONT_RELATIONSHIP_TYPE)
    relationship.set("Target", f"fonts/{DOCX_CJK_FONT_NAME}.odttf")
    entries[relationships_path] = etree.tostring(
        relationships, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    content_types = etree.fromstring(entries["[Content_Types].xml"])
    has_font_type = content_types.xpath(
        "./ct:Default[@Extension='odttf']", namespaces={"ct": CONTENT_TYPES_NS}
    )
    if not has_font_type:
        default = etree.SubElement(content_types, f"{{{CONTENT_TYPES_NS}}}Default")
        default.set("Extension", "odttf")
        default.set(
            "ContentType", "application/vnd.openxmlformats-officedocument.obfuscatedFont"
        )
    entries["[Content_Types].xml"] = etree.tostring(
        content_types, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    entries[font_part] = obfuscate_openxml_font(font_data, font_key)

    temporary = path.with_suffix(".embedded.docx")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as package:
        for name, data in entries.items():
            package.writestr(name, data)
    temporary.replace(path)


def create_fixture_image(path: Path) -> None:
    image = Image.new("RGB", (960, 420), "#EEF3ED")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((48, 48, 912, 372), radius=34, fill="#FFFFFF", outline="#3978C5", width=6)
    draw.rectangle((110, 115, 285, 305), fill="#3978C5")
    draw.rectangle((330, 125, 840, 155), fill="#181A1F")
    draw.rectangle((330, 190, 760, 215), fill="#687076")
    draw.rectangle((330, 248, 820, 273), fill="#A7B7C6")
    font = ImageFont.load_default()
    draw.text((110, 325), "PDF -> DOCX / DOCX -> PDF", fill="#181A1F", font=font)
    image.save(path, format="PNG")


def add_title_block(doc: Document, subtitle: str) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Zero File Conversion Corpus")
    set_run_font(run, size=24, bold=True)
    run.font.color.rgb = BLUE
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(18)
    set_run_font(sub.add_run(subtitle), size=11)


def build_rich_docx(path: Path, image_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc, "Editable layout, CJK text, tables, images, columns, headers and footers")
    intro = doc.add_paragraph()
    set_run_font(intro.add_run("English paragraph. 中文段落用于验证字体、换行与阅读顺序。"))
    doc.add_heading("Document structure", level=1)
    for text in (
        "Preserve headings and paragraph order.",
        "Keep editable text instead of rasterizing the page.",
        "Retain explicit page breaks and repeated document furniture.",
    ):
        paragraph = doc.add_paragraph(style="List Bullet")
        set_run_font(paragraph.add_run(text))

    doc.add_heading("Comparison table", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    values = (
        ("Format", "Expected target", "Important structure"),
        ("PDF", "DOCX", "Reading order and editable text"),
        ("DOCX", "PDF", "Pagination and embedded image"),
        ("中文", "双向", "字体与表格布局"),
    )
    for row, row_values in zip(table.rows, values):
        for cell, value in zip(row.cells, row_values):
            cell.text = ""
            set_run_font(cell.paragraphs[0].add_run(value), size=10, bold=row is table.rows[0])
    set_table_geometry(table, [1.35, 1.65, 3.5])

    doc.add_heading("Embedded image", level=2)
    picture = doc.add_picture(str(image_path), width=Inches(5.9))
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER

    columns = doc.add_section(WD_SECTION.CONTINUOUS)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")
    columns._sectPr.append(cols)
    for index in range(8):
        paragraph = doc.add_paragraph()
        set_run_font(
            paragraph.add_run(
                f"Column sample {index + 1}. Local conversion should keep this reading order and paragraph boundary."
            ),
            size=10,
        )

    one_column = doc.add_section(WD_SECTION.NEW_PAGE)
    single_cols = OxmlElement("w:cols")
    single_cols.set(qn("w:num"), "1")
    one_column._sectPr.append(single_cols)
    doc.add_heading("Explicit page break", level=1)
    paragraph = doc.add_paragraph()
    set_run_font(paragraph.add_run("This section begins on a new page. 第二页内容应保持可编辑。"))
    doc.save(path)


def build_large_docx(path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc, "Twenty-four reproducible pages for timeout and resource tests")
    for page in range(1, 25):
        doc.add_heading(f"Page workload {page}", level=1)
        for paragraph_index in range(6):
            paragraph = doc.add_paragraph()
            set_run_font(
                paragraph.add_run(
                    f"Page {page}, paragraph {paragraph_index + 1}. "
                    "This generated content exercises pagination, repeated text extraction, and bounded conversion resources. "
                    "本段为可重复的中文测试内容。"
                )
            )
        if page != 24:
            doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
    doc.save(path)


def pdf_styles():
    if PDF_FONT_PATH is None:
        raise RuntimeError("CJK fixture font was not configured")
    if PDF_FONT_NAME not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, str(PDF_FONT_PATH)))
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "FixtureTitle",
            parent=base["Title"],
            fontName=PDF_FONT_NAME,
            fontSize=22,
            leading=27,
            textColor=colors.HexColor("#2E74B5"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "heading": ParagraphStyle(
            "FixtureHeading",
            parent=base["Heading1"],
            fontName=PDF_FONT_NAME,
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "FixtureBody",
            parent=base["BodyText"],
            fontName=PDF_FONT_NAME,
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor("#181A1F"),
            spaceAfter=7,
        ),
    }


def build_rich_pdf(path: Path, image_path: Path) -> None:
    styles = pdf_styles()
    doc = BaseDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title="Zero File Conversion Corpus",
        author="Zero test fixture generator",
    )
    page_width, page_height = letter
    gap = 0.25 * inch
    column_width = (page_width - 1.5 * inch - gap) / 2
    frames = [
        Frame(0.75 * inch, 0.7 * inch, column_width, page_height - 1.4 * inch, id="left"),
        Frame(0.75 * inch + column_width + gap, 0.7 * inch, column_width, page_height - 1.4 * inch, id="right"),
    ]
    doc.addPageTemplates(PageTemplate(id="two-column", frames=frames))
    story = [
        Paragraph("Zero File Conversion Corpus", styles["title"]),
        Paragraph("English and 中文 text with two-column reading order.", styles["body"]),
        Paragraph("Editable content", styles["heading"]),
    ]
    for index in range(8):
        story.append(
            Paragraph(
                f"Column paragraph {index + 1}. Preserve paragraph boundaries, punctuation, and reading order. 中文内容 {index + 1}。",
                styles["body"],
            )
        )
    data = [
        ["Format", "Target", "Structure"],
        ["PDF", "DOCX", "Text order"],
        ["DOCX", "PDF", "Pagination"],
        ["中文", "双向", "字体"],
    ]
    table = Table(data, colWidths=[0.8 * inch, 0.8 * inch, 1.35 * inch], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), PDF_FONT_NAME),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#A7B7C6")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend(
        [
            Paragraph("Table and image", styles["heading"]),
            table,
            Spacer(1, 12),
            ReportLabImage(str(image_path), width=2.9 * inch, height=1.27 * inch),
            PageBreak(),
            Paragraph("Explicit page break", styles["heading"]),
            Paragraph("This page verifies page boundaries and headers. 第二页用于测试分页。", styles["body"]),
        ]
    )
    doc.build(story)


def build_scan_pdf(path: Path, image_path: Path) -> None:
    image = Image.open(image_path).convert("RGB")
    canvas = Image.new("RGB", (1275, 1650), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((110, 100), "SCANNED IMAGE-ONLY PDF", fill="#181A1F", font=ImageFont.load_default())
    canvas.paste(image.resize((1055, 462)), (110, 250))
    draw.text((110, 780), "No PDF text objects are present on this page.", fill="#181A1F", font=ImageFont.load_default())
    canvas.save(path, format="PDF", resolution=150)


def build_encrypted_pdf(path: Path) -> None:
    plain = OUTPUT / ".encrypted-source.pdf"
    styles = pdf_styles()
    doc = SimpleDocTemplate(str(plain), pagesize=letter, title="Encrypted fixture")
    doc.build(
        [
            Paragraph("Password-protected conversion fixture", styles["title"]),
            Paragraph("The test password is zero-test and is not user data.", styles["body"]),
        ]
    )
    reader = PdfReader(str(plain))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Title": "Encrypted conversion fixture",
            "/Author": "Zero test fixture generator",
        }
    )
    writer.encrypt("zero-test")
    with path.open("wb") as output:
        writer.write(output)
    plain.unlink()


def build_large_pdf(path: Path) -> None:
    styles = pdf_styles()
    story = []
    for page in range(1, 25):
        story.append(Paragraph(f"Large conversion page {page}", styles["heading"]))
        for paragraph_index in range(10):
            story.append(
                Paragraph(
                    f"Page {page}, paragraph {paragraph_index + 1}. Reproducible workload for resource, timeout, and pagination checks. 中文测试内容。",
                    styles["body"],
                )
            )
        if page != 24:
            story.append(PageBreak())
    doc = SimpleDocTemplate(str(path), pagesize=letter, title="Large conversion fixture")
    doc.build(story)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(65536), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_manifest() -> None:
    cases = [
        ("rich-layout.docx", "docxToPdf", "valid", "webRenderedPdf", ["latin", "cjk", "fonts", "headings", "table", "image", "columns", "header", "footer", "pageBreak"]),
        ("large-structured.docx", "docxToPdf", "valid", "webRenderedPdf", ["large", "pagination", "wordPagination", "cjk", "fonts"]),
        ("malformed.docx", None, "invalidInput", None, ["malformedContainer"]),
        ("~$office-lock.docx", None, "invalidInput", None, ["officeLockFile"]),
        ("rich-layout.pdf", "pdfToDocx", "valid", "layoutPreserving", ["latin", "cjk", "table", "image", "columns", "pageBreak"]),
        ("image-only-scan.pdf", "pdfToDocx", "valid", "layoutPreserving", ["scanned", "imageOnly"]),
        ("encrypted.pdf", "pdfToDocx", "passwordRequired", None, ["encrypted"]),
        ("large-structured.pdf", "pdfToDocx", "valid", "editableReconstruction", ["large", "pagination", "cjk"]),
        ("malformed.pdf", None, "invalidInput", None, ["malformedHeader"]),
        ("unsupported.txt", None, "unsupportedFormat", None, ["unsupportedExtension"]),
    ]
    manifest = {
        "schemaVersion": 1,
        "generatedBy": "scripts/generate-file-conversion-fixtures.py",
        "containsUserData": False,
        "cases": [
            {
                "file": name,
                "direction": direction,
                "expectedPreflight": expected,
                **({"expectedQualityProfile": profile} if profile else {}),
                "coverage": coverage,
                "sha256": sha256(OUTPUT / name),
            }
            for name, direction, expected, profile, coverage in cases
        ],
    }
    (OUTPUT / "expected.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


def verify_docx(path: Path) -> None:
    with zipfile.ZipFile(path) as package:
        names = set(package.namelist())
    required = {"[Content_Types].xml", "word/document.xml"}
    if not required.issubset(names):
        raise RuntimeError(f"{path.name} is missing required DOCX entries")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--cjk-font",
        type=Path,
        default=Path(os.environ["ZERO_FILE_FIXTURE_CJK_FONT"])
        if "ZERO_FILE_FIXTURE_CJK_FONT" in os.environ
        else None,
        help="Path to an embeddable CJK TrueType font, such as OFL Noto Sans SC.",
    )
    return parser.parse_args()


def main() -> None:
    global PDF_FONT_PATH
    args = parse_args()
    if args.cjk_font is None or not args.cjk_font.is_file():
        raise SystemExit(
            "--cjk-font must name an existing embeddable TrueType font; "
            "the generator never downloads fonts implicitly"
        )
    PDF_FONT_PATH = args.cjk_font.resolve()
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for name in FIXTURE_NAMES:
        target = OUTPUT / name
        if target.exists():
            target.unlink()

    image_path = OUTPUT / ".fixture-image.png"
    docx_font_path = OUTPUT / ".fixture-cjk-subset.ttf"
    subset_cjk_font(PDF_FONT_PATH, docx_font_path)
    create_fixture_image(image_path)
    build_rich_docx(OUTPUT / "rich-layout.docx", image_path)
    build_large_docx(OUTPUT / "large-structured.docx")
    embed_docx_font(OUTPUT / "rich-layout.docx", docx_font_path)
    embed_docx_font(OUTPUT / "large-structured.docx", docx_font_path)
    (OUTPUT / "malformed.docx").write_bytes(b"not-a-zip-package")
    (OUTPUT / "~$office-lock.docx").write_bytes(b"office-lock-fixture")
    build_rich_pdf(OUTPUT / "rich-layout.pdf", image_path)
    build_scan_pdf(OUTPUT / "image-only-scan.pdf", image_path)
    build_encrypted_pdf(OUTPUT / "encrypted.pdf")
    build_large_pdf(OUTPUT / "large-structured.pdf")
    (OUTPUT / "malformed.pdf").write_bytes(b"not-a-pdf-header")
    (OUTPUT / "unsupported.txt").write_text("unsupported fixture\n", encoding="utf-8")
    image_path.unlink()
    docx_font_path.unlink()

    verify_docx(OUTPUT / "rich-layout.docx")
    verify_docx(OUTPUT / "large-structured.docx")
    write_manifest()
    print(f"Generated {len(FIXTURE_NAMES) - 1} fixtures in {OUTPUT}")


if __name__ == "__main__":
    main()
